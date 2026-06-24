"""
CV generator — ports the career-ops docx generation to Python.
Uses python-docx to create Calibri/black/US-Letter CVs tailored per job.
"""
import os, re, json
from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List, Optional

OUTPUT_DIR = Path(os.getenv("CV_OUTPUT_DIR", "/tmp/cvs"))
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

CV_DATA = {
    "name": "ANDREA ESTRELLA",
    "contact": "aestrellaoliva@gmail.com  ·  (779) 390-4034  ·  linkedin.com/in/andreaa-estrellaa  ·  github.com/andreastaar  ·  Chicago, IL",
    "education": [
        {
            "school": "Lewis University",
            "location": "Chicago, IL",
            "degree": "M.S. Business Analytics  |  GPA: 4.0/4.0",
            "dates": "Expected May 2027",
            "details": [
                "Currently enrolled graduate student at accredited college/university · Seeking Summer 2026 internship opportunities for masters students",
                "Microsoft Certified: Fabric Analytics Engineer Associate",
                "Selected Coursework: Machine Learning, Statistical Modeling, Data Mining, Time Series Analysis, Linear Algebra, Mathematics, Business Intelligence, Optimization & Decision Theory, Stochastic Modeling, Causal Inference, Economics",
            ]
        },
        {
            "school": "Esan University",
            "location": "Lima, Peru",
            "degree": "B.S. Industrial Engineering  |  GPA: 3.8/4.0 · Graduated with Honors",
            "dates": "Jun 2020",
            "details": [
                "Majors: Data Science & Systems Engineering · Selected Coursework: Deep Learning, Reinforcement Learning, Operations Research, Statistical Analysis, Quantitative Methods",
            ]
        },
    ],
    "experience": [
        {
            "company": "European Dynamics",
            "location": "Greece",
            "role": "AI & Data Project Manager",
            "dates": "Mar 2024 – Jul 2025",
            "bullets": [
                "Led GTM analytics and revenue metrics pipelines; achieved 30% reduction in manual processing through end-to-end execution of data reporting workflows",
                "Validated data inputs, analyzed sales performance and renewal trends; presented clear, actionable findings to business stakeholders; comfortable working with complex datasets spanning millions of rows",
                "Translated complex data outputs into digestible narratives for multiple levels of leadership; problem-solving with incomplete information; communicated recommendations to cross-functional stakeholders aligned with pipeline strategy",
            ]
        },
        {
            "company": "Ransa",
            "location": "Peru",
            "role": "Data & Analytics Project Manager",
            "dates": "Jan 2022 – Feb 2024",
            "bullets": [
                "Led data migration across 6 international accounts alongside data scientists and BI engineers; built and refreshed reporting templates supporting weekly business cadence and stakeholder planning conversations",
                "Architected Power BI + Birst business intelligence dashboards — 10+ dashboards reducing decision-making time 20%; drove reporting efficiency improvements and process automation cutting pipeline time 25%",
            ]
        },
        {
            "company": "Ransa",
            "location": "Peru",
            "role": "Data Science Product Owner",
            "dates": "Aug 2020 – Jan 2022",
            "bullets": [
                "Achieved $100K annual cost savings; built predictive models using Google Cloud, SQL, Hadoop, and Spark for revenue forecasting, demand forecasting, and sales pipeline health analysis",
                "Led 5+ projects delivering data preparation, ad-hoc analysis, and KPI reporting to senior leadership; prior experience in B2B SaaS analytics and go-to-market concepts",
            ]
        },
        {
            "company": "Ferreycorp – Unimaq",
            "location": "Peru",
            "role": "Data & Process Improvement Analyst",
            "dates": "Dec 2018 – Apr 2020",
            "bullets": [
                "Contributed to SAP S4/HANA implementation; led process improvements across warehouse, marketing, and finance — 20% efficiency gain; ISO 9001 recertification 2019 & 2020",
            ]
        },
    ],
    "projects": [
        {
            "name": "RAG Supply Chain Q&A",
            "tech": "LangChain · GPT-4o · FAISS · Streamlit · NLP",
            "link": "github.com/andreastaar/rag-supply-chain-qa",
            "bullet": "End-to-end RAG pipeline; semantic chunking + FAISS vector store; benchmarking framework — 87% answer accuracy, 94% source citation rate, 2.1s avg response",
        },
        {
            "name": "LLM Data Insights Agent",
            "tech": "LangChain Agents · GPT-4o · Pandas · Streamlit",
            "link": "github.com/andreastaar/llm-data-insights",
            "bullet": "Multi-step LLM agent with tool use; compare models for company goals; rubric-based model evaluation; auto-profiles schema, returns chart + written summary",
        },
        {
            "name": "Multilingual LLM Translator API",
            "tech": "OpenAI · FastAPI · Streamlit · NLP",
            "link": "github.com/andreastaar/multilingual-llm-translator",
            "bullet": "Context-aware translation API (EN/ES/PT/FR/DE/JA) with domain modes; confidence scoring; outperforms generic MT on domain-specific terminology",
        },
        {
            "name": "Supply Chain Demand Forecasting",
            "tech": "XGBoost · Scikit-learn · Pandas",
            "link": "github.com/andreastaar/supply-chain-demand-forecasting",
            "bullet": "SKU-level demand forecasting across 5 DCs; systematic feature engineering; 57% MAPE reduction vs naive baseline (18.4% → 7.8%)",
        },
    ],
    "skills": {
        "Languages & Fundamentals": "Python, SQL, Java, JavaScript, C/C++, MATLAB, Ruby, Git, Bash, scripting, programming",
        "ML / AI & LLMs": "PyTorch, TensorFlow, Scikit-learn, XGBoost, LangChain, OpenAI APIs, Hugging Face, RAG, LLMs (GPT-4o, Claude, Mistral, LLaMA, Gemini), NLP, natural language processing, deep learning, neural networks, reinforcement learning, model evaluation, A/B testing, experiments, benchmarking, experimental design",
        "Research & Methods": "Statistics, statistical modeling, econometrics, causal inference, hypothesis testing, A/B experimentation, analytical skills, frameworks, product analytics, product development insights, cross-functional communication, algorithm design, business impact analysis, problem-solving",
        "MLOps & Data Engineering": "Pandas, NumPy, Spark, Hadoop, MapReduce, Hive, ETL pipelines, analytics engineering, AI-ready data foundations, data preparation, data visualization, FastAPI, Streamlit, Docker, Power BI, Microsoft Fabric, Tableau, Excel",
        "Cloud & Platforms": "AWS, Azure, Google Cloud, Snowflake, SAP S4/HANA, Infor WMS, Salesforce, MySQL, SQL Server",
        "Additional": "RevOps analytics · Excel · Google Sheets · Power BI · Tableau · Microsoft Office suite · Alteryx Designer · SaaS/B2B analytics · GTM reporting · Revenue operations · AI fluency · enhance productivity · team-oriented · intellectually curious · proactive · fast-paced · Spanish (fluent) · Portuguese (fluent) · English (fluent) · CFA Society Chicago Member · Lean Six Sigma · Agile/Scrum",
    }
}


def _set_font(run, size_pt: float, bold=False, italic=False, color="000000"):
    run.font.name = "Calibri"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    # Force Calibri via XML
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rPr.insert(0, rFonts)


def _add_section_title(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text.upper())
    _set_font(run, 11, bold=True)
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _add_two_col(doc, left: str, right: str, left_size=11, left_bold=True, right_size=9):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    tab_stops = p.paragraph_format.tab_stops
    # Right-align tab at content width
    tab_stops.add_tab_stop(Inches(7.5), WD_ALIGN_PARAGRAPH.RIGHT)
    run_l = p.add_run(left)
    _set_font(run_l, left_size, bold=left_bold)
    run_tab = p.add_run('\t')
    _set_font(run_tab, right_size)
    run_r = p.add_run(right)
    _set_font(run_r, right_size)
    return p


def _add_bullet(doc, text: str):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    _set_font(run, 9)
    return p


def _add_detail(doc, text: str, italic=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    _set_font(run, 9, italic=italic)
    return p


def inject_keywords(text: str, keywords: List[str]) -> str:
    """Naturally inject missing keywords into text if not present."""
    lower = text.lower()
    additions = [k for k in keywords if k.lower() not in lower]
    if additions:
        text = text + " · " + " · ".join(additions[:5])
    return text


def generate_cv(company: str, role: str, jd: str = "", extra_keywords: Optional[List[str]] = None) -> dict:
    """
    Generate a tailored Word CV for a specific company/role.
    Returns dict with docx_path, pdf_path (pdf generated separately).
    """
    slug = re.sub(r'[^a-z0-9]', '-', company.lower())[:30]
    today = date.today().isoformat()
    filename = f"cv-andrea-{slug}-{today}.docx"
    out_path = OUTPUT_DIR / filename

    doc = Document()

    # Page setup: US Letter, narrow margins
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin    = Inches(0.2)
    section.bottom_margin = Inches(0.2)
    section.left_margin   = Inches(0.3)
    section.right_margin  = Inches(0.3)

    # Remove default styles spacing
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)
    doc.styles['Normal'].paragraph_format.space_after  = Pt(0)

    # ── NAME ──────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    run = p.add_run(CV_DATA["name"])
    _set_font(run, 18, bold=True)
    p.paragraph_format.space_after = Pt(1)

    # ── CONTACT ───────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    run = p.add_run(CV_DATA["contact"])
    _set_font(run, 9)
    p.paragraph_format.space_after = Pt(3)

    # ── EDUCATION ─────────────────────────────────────────────────────────────
    _add_section_title(doc, "Education")
    for edu in CV_DATA["education"]:
        _add_two_col(doc, edu["school"], edu["location"], left_size=11, left_bold=True, right_size=9)
        _add_two_col(doc, edu["degree"], edu["dates"], left_size=10, left_bold=False, right_size=9)
        for detail in edu["details"]:
            _add_detail(doc, detail)

    # ── WORK EXPERIENCE ───────────────────────────────────────────────────────
    _add_section_title(doc, "Work Experience")
    for exp in CV_DATA["experience"]:
        _add_two_col(doc, exp["company"], exp["location"], left_size=11, left_bold=True)
        _add_two_col(doc, exp["role"], exp["dates"], left_size=10, left_bold=False)
        for b in exp["bullets"]:
            _add_bullet(doc, b)

    # ── SELECTED PROJECTS ─────────────────────────────────────────────────────
    _add_section_title(doc, "Selected Projects")
    for proj in CV_DATA["projects"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(0)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(7.5), WD_ALIGN_PARAGRAPH.RIGHT)
        r1 = p.add_run(proj["name"])
        _set_font(r1, 9, bold=True)
        r2 = p.add_run(" | ")
        _set_font(r2, 9)
        r3 = p.add_run(proj["tech"])
        _set_font(r3, 9, italic=True)
        r4 = p.add_run("\t")
        _set_font(r4, 9)
        r5 = p.add_run(proj["link"])
        _set_font(r5, 9)
        _add_bullet(doc, proj["bullet"])

    # ── SKILLS ────────────────────────────────────────────────────────────────
    _add_section_title(doc, "Skills")
    skills = dict(CV_DATA["skills"])

    # Inject extra keywords into Additional if provided
    if extra_keywords:
        skills["Additional"] = inject_keywords(skills["Additional"], extra_keywords)

    for label, content in skills.items():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run_l = p.add_run(label + ": ")
        _set_font(run_l, 9, bold=True)
        run_r = p.add_run(content)
        _set_font(run_r, 9)

    doc.save(str(out_path))
    return {
        "docx_path": str(out_path),
        "pdf_path":  str(out_path).replace(".docx", ".pdf"),
        "filename":  filename,
    }
